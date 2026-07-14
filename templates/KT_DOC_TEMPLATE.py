# -*- coding: utf-8 -*-
"""
KT Word-document generator template.

HOW TO USE THIS TEMPLATE (for a Claude agent in a new project):
  1. Research the target project first (notebooks/scripts, existing reports, output files) —
     see README_TEMPLATE_PROMPT.md in this same templates/ folder for the research approach.
     Verify every claim against source code and result files before writing it down here.
  2. Fill in the CONFIG block below and the content lists in each section (models, glossary
     terms, notebook tiers, pending work, etc.) with real, verified content for the target
     project. Everything marked {{PLACEHOLDER}} or "TODO" must be replaced.
  3. Run this script from the project root: `python KT_DOC_TEMPLATE.py`. It writes a .docx
     next to itself. Rename the output file appropriately and move it to the project root.
  4. Do NOT invent numbers, filenames, or function names — every fact in this document should
     be traceable to something you actually read in the target repo. If you can't verify a
     claim from an original report/prior KT doc, mark it as unverified rather than restating it.
  5. Delete this docstring block before generating the final document (optional, but tidy).

This template intentionally mirrors the structure that worked well for a prior project:
Title page -> Executive Summary -> Project Overview -> Glossary -> Pipeline -> Approaches/Models
-> Key Results -> Reading Guide (tiered) -> Environment Setup -> Pending Work -> Quick Reference.
Feel free to drop sections that don't apply, but keep the "verify + flag caveats" spirit.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================================
# CONFIG — fill in for the target project
# ============================================================================
PROJECT_TITLE      = '{{Project Name}}'                     # e.g. 'HAR-RS Model & Extensions'
PROJECT_SUBTITLE   = 'Knowledge Transfer Document'
PROJECT_TAGLINE     = '{{One-line domain description}}'      # e.g. 'Volatility Forecasting Research — NIFTY / India VIX'
PREPARED_BY         = '{{Your name}}'
PREPARED_DATE       = '{{Month Year}}'

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Colour palette — change to taste, keep good contrast for header/body/accent
DARK_BLUE  = RGBColor(0x1F, 0x3A, 0x5F)
MID_BLUE   = RGBColor(0x2E, 0x6D, 0xA4)
GREEN      = RGBColor(0x19, 0x7A, 0x3A)
RED        = RGBColor(0xC0, 0x39, 0x2B)
PURPLE     = RGBColor(0x8E, 0x44, 0xAD)
GREY       = RGBColor(0x7F, 0x8C, 0x8D)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# ============================================================================
# Style helpers — reusable, do not need to change
# ============================================================================

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def shade_para(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def add_heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color is None:
        color = DARK_BLUE if level == 1 else MID_BLUE
    for run in p.runs:
        run.font.color.rgb = color
    return p

def add_para(text, bold=False, color=None, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def add_note(label, text, fill='E8F4FD', label_color=None):
    """Callout box. Use fill='E9F7EF'/label_color=GREEN for a positive/winner note,
    fill='FDEDEC'/label_color=RED for a warning/caveat, fill='FEF9E7' for a neutral note."""
    p = doc.add_paragraph()
    shade_para(p, fill)
    lbl = p.add_run(label + '  ')
    lbl.font.bold = True
    lbl.font.size = Pt(10)
    if label_color:
        lbl.font.color.rgb = label_color
    body = p.add_run(text)
    body.font.size = Pt(10)
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, '1F3A5F')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        fill = 'F7FAFD' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            shade_cell(cell, fill)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def add_labelled_block(title, rows, fill='EBF5FB', name_color='1F3A5F'):
    """One 'card' for a model/approach/experiment: title bar + label:value rows."""
    p_title = doc.add_paragraph()
    shade_para(p_title, fill)
    r = p_title.add_run(title)
    r.font.bold = True
    r.font.size = Pt(11)
    c = name_color
    r.font.color.rgb = RGBColor(int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16))
    p_title.paragraph_format.space_before = Pt(8)
    p_title.paragraph_format.left_indent  = Cm(0.3)

    for label, val in rows:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.6)
        p.paragraph_format.space_before = Pt(1)
        lbl = p.add_run(label + ': ')
        lbl.font.bold = True
        lbl.font.size = Pt(9.5)
        body = p.add_run(val)
        body.font.size = Pt(9.5)
    doc.add_paragraph()

def add_tier_header(text, fill_hex):
    p = doc.add_paragraph()
    shade_para(p, fill_hex)
    r = p.add_run('  ' + text)
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = WHITE
    p.paragraph_format.space_before = Pt(6)
    return p

def add_file_entry(name, subtitle, desc, name_color=MID_BLUE):
    p_nb = doc.add_paragraph()
    p_nb.paragraph_format.left_indent  = Cm(0.4)
    p_nb.paragraph_format.space_before = Pt(6)
    r1 = p_nb.add_run(name + '\n')
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = name_color
    r2 = p_nb.add_run(subtitle + '\n')
    r2.font.bold = True
    r2.font.size = Pt(9)
    r3 = p_nb.add_run(desc)
    r3.font.size = Pt(9)

# ============================================================================
# 0. TITLE PAGE
# ============================================================================
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run(PROJECT_TITLE)
tr.font.size = Pt(26)
tr.font.bold = True
tr.font.color.rgb = DARK_BLUE

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run(PROJECT_SUBTITLE)
sr.font.size = Pt(16)
sr.font.color.rgb = MID_BLUE

doc.add_paragraph()

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = meta_p.add_run(PROJECT_TAGLINE + '\n')
r1.font.size = Pt(11)
r2 = meta_p.add_run(f'Prepared by: {PREPARED_BY}   |   {PREPARED_DATE}')
r2.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph()

add_note(
    'Purpose of this document:',
    'Written for the researcher, engineer, or PM who will continue or review this work. '
    'It explains what was researched, what was found, what was tried and rejected, and — most '
    'importantly — exactly which files to open and in what order, so you never have to guess. '
    'Every claim below has been checked against the actual code and result files, not just '
    'against prose reports — anywhere that was not possible, it is flagged explicitly.',
    fill='EBF5FB'
)

doc.add_page_break()

# ============================================================================
# 1. EXECUTIVE SUMMARY
# ============================================================================
add_heading('1. Executive Summary')

add_para(
    '{{TODO: 2-4 sentences — what question this project answers, and the overall approach '
    'taken (method/model family/framework used).}}'
)
add_para(
    '{{TODO: 1-2 sentences on how many approaches were compared and what evaluation method/metric '
    'was used to compare them.}}'
)

add_note(
    'Winner:',
    '{{TODO: name the best-performing approach and its headline metric(s). If you found during '
    'verification that this result cannot be fully reproduced or trusted (e.g. missing source '
    'code, unreproducible numbers, outlier-distorted stats), say so in the SAME note, not just '
    'later in the document — this is the single most important thing a successor needs to know '
    'before relying on this result.}}',
    fill='E9F7EF', label_color=GREEN
)

doc.add_page_break()

# ============================================================================
# 2. PROJECT OVERVIEW
# ============================================================================
add_heading('2. Project Overview')

add_heading('2.1 Objective', level=2)
add_para('{{TODO: what is this project trying to achieve, and why (business/research motivation)?}}')

add_heading('2.2 Scope & Data Context', level=2)
add_bullet('{{TODO: universe / subject of study}}')
add_bullet('{{TODO: data frequency, time range, size}}')
add_bullet('{{TODO: key data files/sources}}')
add_bullet('{{TODO: evaluation horizons/targets, if applicable}}')
add_bullet('{{TODO: train/test or in-sample/out-of-sample split description}}')

add_heading('2.3 Repository Structure', level=2)
add_table(
    ['Folder / File', 'Contents'],
    [
        ['{{data/}}',            '{{TODO}}'],
        ['{{notebooks/ or src/}}', '{{TODO — how many files, what kind of work}}'],
        ['{{reports/ or outputs/}}', '{{TODO}}'],
        ['{{shared utility module}}', '{{TODO}}'],
        ['{{requirements/env file}}', '{{TODO}}'],
        ['{{final written report, if any}}', '{{TODO}}'],
        # add/remove rows to match the real repo
    ],
    col_widths=[5.5, 11]
)

doc.add_page_break()

# ============================================================================
# 3. CORE CONCEPTS GLOSSARY
# ============================================================================
add_heading('3. Core Concepts Glossary')
add_para('Quick reference for anyone unfamiliar with the terminology used throughout the research.')
doc.add_paragraph()

add_table(
    ['Term', 'Definition'],
    [
        # TODO: replace with the actual domain glossary for the target project.
        ['{{Term 1}}', '{{Definition — plain language, tie back to how it is computed/used here.}}'],
        ['{{Term 2}}', '{{...}}'],
        ['{{Metric used for evaluation}}', '{{Formula/definition. Note explicitly if there is a '
         'similarly-named-but-different metric elsewhere in the codebase (a common source of '
         'confusion worth flagging up front).}}'],
        ['{{Key methodological safeguard, e.g. "Look-ahead bias"}}',
         '{{What it means and how the pipeline avoids it.}}'],
    ],
    col_widths=[4.5, 12]
)

doc.add_page_break()

# ============================================================================
# 4. PIPELINE / METHODOLOGY
# ============================================================================
add_heading('4. Data Pipeline & Methodology')
add_para(
    '{{TODO: one sentence on where the pipeline logic lives (shared module vs. per-notebook) '
    'and any overarching design principle (e.g. strictly causal / no look-ahead).}}'
)

add_heading('4.1 Step-by-Step Pipeline', level=2)
steps = [
    # (Step title, detail — cite the actual function/file for each step)
    ('Step 1 — {{Load raw data}}',       '{{module.function()}}\n{{what it does}}'),
    ('Step 2 — {{Transform/derive core series}}', '{{module.function()}}\n{{what it does}}'),
    ('Step 3 — {{Feature/label construction}}', '{{...}}'),
    ('Step 4 — {{Preprocessing / normalization}}', '{{...}}'),
    ('Step 5 — {{Fit & forecast/predict}}', '{{...}}'),
    ('Step 6 — {{Evaluate}}', '{{Metrics used, validation scheme}}'),
]
for title, detail in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    r1 = p.add_run(title + '\n')
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = MID_BLUE
    r2 = p.add_run(detail)
    r2.font.size = Pt(9.5)

add_heading('4.2 Key Shared Functions', level=2)
add_table(
    ['Function', 'File', 'What it does'],
    [
        # TODO: list actual shared utility functions, verified by reading the source
        ['{{func_name()}}', '{{file path}}', '{{one-line description}}'],
    ],
    col_widths=[5, 3.8, 7.7]
)

doc.add_page_break()

# ============================================================================
# 5. APPROACHES / MODELS IMPLEMENTED
# ============================================================================
add_heading('5. Approaches Implemented')
add_para(
    '{{TODO: how many approaches were tried, what estimation/validation method was used '
    'across all of them.}}'
)
doc.add_paragraph()

# One add_labelled_block() call per approach. Colors: blue=baseline, purple=extension,
# green=winner, red=rejected. IMPORTANT: only list a notebook/file as "Primary implementation"
# if you have actually verified the code exists there — if you can only find the RESULT
# (e.g. in an output spreadsheet) but not the source code that produced it, say so explicitly
# in the 'Key finding' or a dedicated row instead of naming a file that doesn't contain it.
add_labelled_block(
    '{{Approach 1 — Baseline}}',
    [
        ('Description',    '{{what it is}}'),
        ('Features/inputs', '{{...}}'),
        ('Key finding',    '{{...}}'),
        ('Primary file',   '{{path, or "Not verified in current repo" if the code cannot be found}}'),
    ],
    fill='EBF5FB', name_color='1F3A5F'
)
add_labelled_block(
    '{{Approach 2 — Extension}}',
    [
        ('Description',    '{{...}}'),
        ('Features/inputs', '{{...}}'),
        ('Key finding',    '{{...}}'),
        ('Primary file',   '{{...}}'),
    ],
    fill='F4ECF7', name_color='8E44AD'
)
add_labelled_block(
    '{{Approach N — Best / Winner}} ★',
    [
        ('Description',    '{{...}}'),
        ('Features/inputs', '{{...}}'),
        ('Key finding',    '{{...}}'),
        ('Primary file',   '{{...}}'),
    ],
    fill='D5F5E3', name_color='1A7A3A'
)
add_labelled_block(
    '{{Rejected approach}}',
    [
        ('Description',    '{{...}}'),
        ('Key finding',    '{{why it was rejected}}'),
        ('Primary file',   '{{...}}'),
    ],
    fill='FDEDEC', name_color='C0392B'
)

doc.add_page_break()

# ============================================================================
# 6. KEY RESULTS
# ============================================================================
add_heading('6. Key Results')

add_heading('6.1 Comparison Summary', level=2)
add_para('{{TODO: one line on what the table below shows and where the raw numbers live.}}', size=9)
doc.add_paragraph()

add_table(
    ['Approach', 'Setting', 'Primary metric', 'Secondary metric', 'Verdict'],
    [
        # TODO: pull real numbers from actual result files, not from memory of the prose report
        ['{{Approach}}', '{{e.g. horizon/split}}', '{{value}}', '{{value}}', '{{Baseline/Best/Rejected}}'],
    ],
    col_widths=[5, 3, 3, 3, 3]
)

add_note(
    'Note:',
    '{{TODO: point to exactly where the full numeric tables live (report file + section/table '
    'numbers, or output file names).}}',
    fill='FEF9E7'
)

add_heading('6.2 Statistical Validation', level=2)
add_para(
    '{{TODO: describe whatever unbiasedness/significance/robustness check was used (e.g. '
    'Mincer-Zarnowitz regression, residual diagnostics, cross-validation stability) and its result '
    'for the winning approach.}}'
)

add_heading('6.3 Limitations', level=2)
add_bullet('{{TODO: known weaknesses of the winning approach (regimes/segments where it fails).}}')
add_bullet('{{TODO: incomplete work, e.g. a broken hyperparameter search — state status precisely, '
           'do not restate unverified parameter values as fact.}}')
add_bullet('{{TODO: any metric-definition traps — e.g. two similarly named metrics that must not '
           'be conflated, or a summary statistic distorted by outliers. Flag these even if the '
           'original report did not.}}')
add_bullet('{{TODO: generalizability caveats — what this was and was not calibrated/tested on.}}')

doc.add_page_break()

# ============================================================================
# 7. FILE / NOTEBOOK READING GUIDE
# ============================================================================
add_heading('7. File Reading Guide')
add_para(
    '{{TODO: total file count and a one-line statement that most are exploratory/superseded — '
    'the reader does not need to open all of them.}}'
)
doc.add_paragraph()

add_tier_header('TIER 1 — Must Read (start here)', '1F3A5F')
add_file_entry(
    '{{file name}}',
    '{{one-line role, e.g. "Comprehensive EDA — read this FIRST"}}',
    '{{TODO: why this file matters and what a reader will learn from it.}}'
)
add_file_entry(
    '{{main implementation file}}',
    '{{one-line role}}',
    '{{TODO: state precisely which approaches this file actually implements — verified by reading '
    'it, not assumed from its name or from what a report claims about it.}}'
)
doc.add_paragraph()

add_tier_header('TIER 2 — Read If Extending the Work', '2E6DA4')
add_file_entry('{{file name}}', '{{role}}', '{{why/when to read it}}')
doc.add_paragraph()

add_tier_header('TIER 3 — Skip (historical / superseded)', '7F8C8D')
add_para(
    'These files are part of the iterative development history — early or rough versions of '
    'ideas later formalised in Tier 1/2 files. Safely skip all of these.',
    size=9
)
doc.add_paragraph()
add_table(
    ['File', 'Why you can skip it'],
    [
        ['{{file name}}', '{{e.g. "superseded by X"}}'],
    ],
    col_widths=[6.5, 10]
)

doc.add_page_break()

# ============================================================================
# 8. ENVIRONMENT SETUP & HOW TO RUN
# ============================================================================
add_heading('8. Environment Setup & How to Run')

add_heading('8.1 Prerequisites', level=2)
add_bullet('{{Python/runtime version}}')
add_bullet('{{Virtual environment location, if any}}')
add_bullet('{{Recommended IDE/tooling}}')

add_heading('8.2 Setup Steps', level=2)
setup_steps = [
    '{{Open the repo at ...}}',
    '{{Activate environment: ...}}',
    '{{Install dependencies: ...}}',
    '{{Launch/entry point: ...}}',
    '{{Open files in order: ... first, then ...}}',
]
for s in setup_steps:
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(s)
    run.font.size = Pt(10)

add_heading('8.3 Data Files Required', level=2)
add_table(
    ['File', 'Location', 'Used for'],
    [
        ['{{file}}', '{{path}}', '{{purpose}}'],
    ],
    col_widths=[5.5, 2.5, 8.5]
)

add_note(
    'Data refresh:',
    '{{TODO: how to extend/refresh the data, any date-stamping convention in filenames, and '
    'where the path is configured (or note if it is hardcoded per-file with no central config).}}',
    fill='FEF9E7'
)

doc.add_page_break()

# ============================================================================
# 9. PENDING WORK
# ============================================================================
add_heading('9. Pending Work & Future Directions')

add_heading('9.1 High Priority (incomplete work)', level=2)
add_table(
    ['Task', 'Status', 'Where to look'],
    [
        ['{{task}}', '{{status, e.g. "broken / all trials failed"}}', '{{file(s)}}'],
    ],
    col_widths=[5.5, 3.5, 7.5]
)

add_heading('9.2 Research Extensions Worth Exploring', level=2)
add_table(
    ['Extension', 'Description'],
    [
        ['{{idea}}', '{{why it might help, and what it would take}}'],
    ],
    col_widths=[4.5, 12]
)

doc.add_page_break()

# ============================================================================
# 10. QUICK REFERENCE CARD
# ============================================================================
add_heading('10. Quick Reference Card')
add_para('One-page cheat sheet for day-to-day reference.')
doc.add_paragraph()

add_table(
    ['Question', 'Answer'],
    [
        ['What is the best approach?',            '{{...}}'],
        ['Where is the main implementation?',      '{{...}}'],
        ['Where is the exploratory analysis?',     '{{...}}'],
        ['Where are shared utility functions?',    '{{...}}'],
        ['Where are output/result files?',         '{{...}}'],
        ['Where is the final written report?',     '{{...}}'],
        ['What evaluation metric was used?',       '{{...}}'],
        ['What validation method?',                '{{...}}'],
        ['How many approaches were compared?',     '{{...}}'],
        ['Is there any known-broken component?',   '{{...}}'],
        ['Which files must I read first?',         '{{...}}'],
        ['Which files can I skip?',                '{{See Tier 3 above}}'],
    ],
    col_widths=[6.5, 10]
)

# ============================================================================
# SAVE
# ============================================================================
out_path = f'{PROJECT_TITLE} KT Document.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
